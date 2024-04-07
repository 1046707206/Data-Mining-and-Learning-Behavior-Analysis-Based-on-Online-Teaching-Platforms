#!/usr/bin/env python3
# coding: utf-8
# File: crime_mining.py
# Author: lhy<lhy_in_blcu@126.com,https://huangyong.github.io>
# Date: 18-7-24
import csv

from py2neo import Graph, Node, Relationship
import pandas as pd
from sentence_parser import *
import re
from collections import Counter
from GraphShow import *
from keywords_textrank import *
import docx

'''事件挖掘'''
class CrimeMining:
    def __init__(self):
        self.textranker = TextRank()
        self.parser = LtpParser()
        self.ners = ['nh', 'ni', 'ns']
        self.ner_dict = {
        'nh':'人物',
        'ni':'机构',
        'ns':'地名'
        }
        self.graph_shower = GraphShow()

    '''移除括号内的信息，去除噪声'''
    def remove_noisy(self, content):
        p1 = re.compile(r'（[^）]*）')
        p2 = re.compile(r'\([^\)]*\)')
        return p2.sub('', p1.sub('', content))

    '''收集命名实体'''
    def collect_ners(self, words, postags):
        ners = []
        for index, pos in enumerate(postags):
            if pos in self.ners:
                ners.append(words[index] + '/' + pos)
        return ners

    '''对文章进行分句处理'''
    def seg_content(self, content):
        return [sentence for sentence in re.split(r'[？?！!。；;：:\n\r]', content) if sentence]

    '''对句子进行分词，词性标注处理'''
    def process_sent(self, sent):
        words, postags = self.parser.basic_process(sent)
        return words, postags

    '''构建实体之间的共现关系'''
    def collect_coexist(self, ner_sents, ners):
        co_list = []
        for sent in ner_sents:
            words = [i[0] + '/' + i[1] for i in zip(sent[0], sent[1])]
            co_ners = set(ners).intersection(set(words))
            co_info = self.combination(list(co_ners))
            co_list += co_info
        if not co_list:
            return []
        return {i[0]:i[1] for i in Counter(co_list).most_common()}

    '''列表全排列'''
    def combination(self, a):
        combines = []
        if len(a) == 0:
            return []
        for i in a:
            for j in a:
                if i == j:
                    continue
                combines.append('@'.join([i, j]))
        return combines

    '''抽取出事件三元组'''
    def extract_triples(self, words, postags):
        svo = []
        tuples, child_dict_list = self.parser.parser_main(words, postags)
        for tuple in tuples:
            rel = tuple[-1]
            if rel in ['SBV']:
                sub_wd = tuple[1]
                verb_wd = tuple[3]
                obj = self.complete_VOB(verb_wd, child_dict_list)
                subj = sub_wd
                verb = verb_wd
                if not obj:
                    svo.append([subj, verb])
                else:
                    svo.append([subj, verb+obj])
        return svo

    '''过滤出与命名实体相关的事件三元组'''
    def filter_triples(self, triples, ners):
        ner_triples = []
        for ner in ners:
            for triple in triples:
                if ner in triple:
                    ner_triples.append(triple)
        return ner_triples

    '''根据SBV找VOB'''
    def complete_VOB(self, verb, child_dict_list):
        for child in child_dict_list:
            wd = child[0]
            attr = child[3]
            if wd == verb:
                if 'VOB' not in attr:
                    continue
                vob = attr['VOB'][0]
                obj = vob[1]
                return obj
        return ''

    '''对文章进行关键词挖掘'''
    def extract_keywords(self, words_list):
        return self.textranker.extract_keywords(words_list, 10)

    '''基于文章关键词，建立起实体与关键词之间的关系'''
    def rel_entity_keyword(self, ners, keyword, subsent):
        events = []
        rels = []
        sents = []
        ners = [i.split('/')[0] for i in set(ners)]
        keyword = [i[0] for i in keyword]
        for sent in subsent:
            tmp = []
            for wd in sent:
                if wd in ners + keyword:
                    tmp.append(wd)
            if len(tmp) > 1:
                sents.append(tmp)
        for ner in ners:
            for sent in sents:
                if ner in sent:
                    tmp = ['->'.join([ner, wd]) for wd in sent if wd in keyword and wd != ner and len(wd) > 1]
                    if tmp:
                        rels += tmp

        for e in set(rels):
            events.append([e.split('->')[0], e.split('->')[1]])
        return events


    '''利用标点符号，将文章进行短句切分处理'''
    def seg_short_content(self, content):
        return [sentence for sentence in re.split(r'[，,？?！!。；;：:\n\r\t ]', content) if sentence]

    '''挖掘主控函数'''
    def main(self, content):
        if not content:
            return []
        # 对文章进行去噪处理
        content = self.remove_noisy(content)
        # 对文章进行长句切分处理
        sents = self.seg_content(content)
        # 对文章进行短句切分处理
        subsents = self.seg_short_content(content)
        subsents_seg = []
        # words_list存储整篇文章的词频信息
        words_list = []
        # ner_sents保存具有命名实体的句子
        ner_sents = []
        # ners保存命名实体
        ners = []
        # triples保存主谓宾短语
        triples = []
        # 存储文章事件
        events = []
        for sent in subsents:
            words, postags = self.process_sent(sent)
            words_list += [[i[0], i[1]] for i in zip(words, postags)]
            subsents_seg.append([i[0] for i in zip(words, postags)])
            ner = self.collect_ners(words, postags)
            if ner:
                triple = self.extract_triples(words, postags)
                if not triple:
                    continue
                triples += triple
                ners += ner
                ner_sents.append([words, postags])

        # 获取文章关键词, 并图谱组织, 这个可以做
        keywords = [i[0] for i in self.extract_keywords(words_list)]
        for keyword in keywords:
            name = keyword
            cate = '关键词'
            events.append([name, cate])
        # 对三元组进行event构建，这个可以做
        for t in triples:
            if (t[0] in keywords or t[1] in keywords) and len(t[0]) > 1 and len(t[1]) > 1:
                events.append([t[0], t[1]])

        # 获取文章词频信息话，并图谱组织，这个可以做
        word_dict = [i for i in Counter([i[0] for i in words_list if i[1][0] in ['n', 'v'] and len(i[0]) > 1]).most_common()][:10]
        for wd in word_dict:
            name = wd[0]
            cate = '高频词'
            events.append([name, cate])

        #　获取全文命名实体，这个可以做
        ner_dict = {i[0]:i[1] for i in Counter(ners).most_common()}
        for ner in ner_dict:
            name = ner.split('/')[0]
            cate = self.ner_dict[ner.split('/')[1]]
            events.append([name, cate])

        # 获取全文命名实体共现信息,构建事件共现网络
        co_dict = self.collect_coexist(ner_sents, list(ner_dict.keys()))
        co_events = [[i.split('@')[0].split('/')[0], i.split('@')[1].split('/')[0]] for i in co_dict]
        events += co_events
        #将关键词与实体进行关系抽取
        events_entity_keyword = self.rel_entity_keyword(ners, keywords, subsents_seg)
        events += events_entity_keyword
        #对事件网络进行图谱化展示
        self.graph_shower.create_page(events)
        # print(events)
        column=['column1','column2']
        test=pd.DataFrame(columns=column,data=events)
        test.to_csv('D:/桌面/课程.csv',encoding='utf_8_sig')
        # df = pd.read_csv('C:/Users/lxl/Desktop/1.csv', error_bad_lines=False, encoding='utf-8')
        # df = df.fillna('unknown')  # 填充缺失值


        # test.to_csv('test.csv',encoding='utf-8')
        #
        #
        #
        # self.g = Graph("http://localhost:7474", auth=("neo4j", ""))
        # 连接neo4j数据库，输入地址、用户名、密码
        # cur_dir = '/'.join(os.path.abspath(__file__).split('/')[:-1])
        # self.data_path = os.path.join(cur_dir, 'data/medical.json')

        # self.graph = Graph("http://localhost:7474", ayth=("neo4j", ""))
        # self.g.delete_all()
        # self.g.begin()

        #
        # with open('test.csv', 'r', encoding='utf-8') as f:
        #     reader = csv.reader(f)
        #     for item in reader:
        #         if reader.line_num == 1:  # 过滤掉 csv 第 1 列， 0， 1， 2
        #             continue
        #         print("当前行数: ", reader.line_num-1, "当前内容: ", item)
        #         start_node = Node("first", name=item[1])
        #         end_node = Node("second", value=item[2])
        #         relation = Relationship(start_node, item[0], end_node)
        #
        #         #  建立结点：
        #         for item in start_node:
        #             shiti, label = item.split()
        #             cypher_ = "CREATE (:" + label + " {name:'" + shiti + "'})     "
        #             self.g.run(cypher_)
        #         #  建立结点：
        #         for item in end_node:
        #             shiti, label = item.split()
        #             cypher_ = "CREATE (:" + label + " {name:'" + shiti + "'})     "
        #             self.g.run(cypher_)
        #         # 建立关系 ：
        #         for item in reader:
        #             cypher_ = "MATCH  (a:" + item[1] + "),(b:" + item[2] + ") WHERE a.name = '" + item[
        #                 1] + "' AND b.name = '" + item[2] + "' CREATE (a)-[r:relation]->(b)"
        #             self.g.run(cypher_)
        #
        #
        #
        #
        #         self.g.merge(start_node, "first", "name")
        #         self.g.merge(end_node, "first", "value")
        #         self.g.merge(relation, "值", "属性")


        # fr=open('test.csv','r',encoding='utf-8')
        # lst=[]
        # node=[]
        # for row in csv.reader(fr):
        #     lst_ = []
        #     lst_.append(row[1])
        #     lst_.append(row[2])
        #     lst.append(lst_)
        #     node.append(row[1] + ' ' + row[2])
        # # graph = Graph('bolt://localhost:7474', name="neo4j", password="")
        # self.g = Graph("http://localhost:7474", auth=("neo4j", ""))
        # for item in lst:
        #     cypher_ = "MATCH (a:" + item[0] + "{name:'" + item[1] + "'}) SET a." + item[0] + " = '" + item[1] + "'"
        #     self.g.run(cypher_)

file=docx.Document("D:\\桌面\\课程.docx")
# content1 = file.paragraphs
# print(file.paragraphs)

content1=""
for para in file.paragraphs:
    print(para.text)
    content1=content1+para.text



# content1 = """
#         中国共产党第一次全国代表大会
#     1921年7月23日至31日在上海举行。出席代表12人，代表党员50多人。会议在最后一天转移到浙江嘉兴南湖的游船上举行。大会通过了中国共产党的第一个纲领。
#     　　1920年夏至1921年春，随着马克思主义在中国的广泛传播，中国工人运动的蓬勃兴起，作为两者结合产物的中国共产党早期组织，在上海、北京、武汉、长沙、济南、广州以及赴日、旅欧留学生中相继成立，建党条件基本成熟，召开全国代表大会也在建党骨干中开始酝酿。
#     　　6月3日，共产国际代表马林取道欧洲来到上海，与从西伯利亚南下的另一位国际代表尼科尔斯基会合。他们很快与陈独秀离沪期间主持上海党组织工作的李达、李汉俊取得联系，并交换了情况。共产国际代表建议及早召开党的代表大会，宣告中国共产党的正式成立。
#     　　李达、李汉俊在征询陈独秀、李大钊的意见并获得同意后，分别写信给各地党组织，要求每个地区派出两位代表到上海出席党的全国代表大会。
#     　　7月中下旬，设在法租界白尔路389号（今太仓路127号）的博文女校，陆续住进了一批教师、学生模样的青年人，以北京大学师生暑期考察团的名义，来上海参加这次历史性的聚会。代表们到齐以后，就便在住处开了预备会。
#     　　7月23日晚，中国共产党第一次全国代表大会在上海法租界望志路106号（今兴业路76号）正式开幕。会址设在李书城、李汉俊兄弟住宅，大家围坐在客厅长餐桌四周，室内没有特别布置，陈设简单，气氛庄重。出席者有上海的李汉俊、李达；北京的张国焘、刘仁静；长沙的毛泽东、何叔衡；武汉的董必武、陈潭秋；济南的王尽美、邓恩铭；广州的陈公博；留日学生周佛海以及陈独秀委派的包惠僧。陈独秀和李大钊因公务在身未出席会议，而在代表们心目中他们仍是党的主要创始人和领袖。
#     　　两位共产国际代表出席了一大开幕会议，并发表热情的讲话。马林首先指出：中国共产党的成立具有重大的世界意义，第三国际增加了一个东方支部，苏俄布尔什维克又多了一个亲密战友，并对中共提出了建议和希望。尼科尔斯基介绍了共产国际远东局的情况，要求中共把工作进程及时报告远东局。
#     　　接着，代表们商讨了会议的任务和议题，一致确定先由各地代表报告本地工作，再讨论并通过党的纲领和今后工作计划，最后选举中央领导机构。
#     　　7月24日举行第二次会议，各地代表报告本地区党团组织的状况和工作进程，并交流了经验体会。25、26日休会，用于起草党的纲领和今后工作计划。27、28和29日三天，分别举行三次会议，集中议论此前起草的纲领和决议。讨论认真热烈，大家各抒己见，既有统一的认识，又在某些问题引起争论，会议未作出决定。
#         7月30日晚，一大举行第六次会议，原定议题是通过党的纲领和决议，选举中央机构。会议刚开始几分钟，法租界巡捕房密探突然闯入，这次会被迫中断。
#     　　一大第六次会议刚开始，就遭受法租界巡捕房的侵扰。首先闯入会场的叫程子卿，他是黄金荣的把兄弟，利用这层关系进入巡捕房，任华人探长。原来马林由莫斯科途经欧洲来华，曾在维也纳被警察局拘捕，虽经营救获释，但其行动一直作为“赤色分子”被严密监视。具有丰富秘密工作经验的马林，警觉地说这人一定是“包打听”，建议立即停会，大家分头离开。
#     　　果然，十几分钟后两辆警车包围了一大会址，法籍警官亲自带人进入室内询问搜查，没有找到多少证据，威胁警告一番后撤走了。这次冲击虽然没有带来重大损失，毕竟一大不能再在原址进行了。转移出来的一大代表当晚集中于李达寓所商讨，大家一致认为会议不能在上海举行了，有人提议到杭州开会，又有的提出杭州过于繁华，容易暴露目标。当时在场的李达夫人王会悟提出：不如到我的家乡嘉兴南湖开会，离上海很近，又易于隐蔽。大家都赞成，觉得这个安排很妥当。
#     　　第二天清晨，代表们分两批乘火车前往嘉兴。两位国际代表目标太大，李汉俊、陈公博也因经历一场虚惊，都未去嘉兴。10时左右，代表们先后到达嘉兴车站，在鸳湖旅馆稍事休息后，登上事先租好的南湖画舫。
#     　　这是一个阴天，下起了蒙蒙细雨，游人渐渐离去，秀丽的南湖显得格外清静优雅。11时许，一大会议在缓缓划行的画舫上开始了。
#     　　南湖会议继续着上海30日未能进行的议题，先讨论并通过《中国共产党的第一个纲领》，这份15条约700字的简短纲领，确定了党的名称、奋斗目标、基本政策、提出了发展党员、建立地方和中央机构等组织制度，兼有党纲和党章的内容，是党的第一个正式文献。
#     　　接着讨论并通过《中国共产党的第一个决议》，对今后党的工作作出安排部署，鉴于党的力量还弱小，决定以主要精力建立工会组织，指导工人运动和做好宣传工作，并要求与其他政党关系上保持独立政策，强调与第三国际建立紧密关系。
#     　　下午5时，天气转晴，湖面上一艘汽艇向画舫急驰而来。大家因有上海的经历而提高了警惕，立即藏起文件，桌上摆出麻将牌，装扮成游客。后来打听到这是当地士绅的私人游艇，大家才松了一口气，会议仍继续进行。
#     　　最后，一大选举中央领导机构，代表们认为目前党员人数少、地方组织尚不健全，暂不成立中央委员会，先建立三人组成的中央局，并选举陈独秀任书记，张国焘为组织主任，李达为宣传主任。党的第一个中央机关由此产生。会议在齐呼“第三国际万岁”“中国共产党万岁”声中闭幕。
#     　　一大召开标志着中国共产党的正式成立，犹如一轮红日在东方冉冉升起，照亮了中国革命的前程。这是近代中国社会进步和革命发展的客观要求，是开天辟地的大事变。自从有了中国共产党，中国革命的面目就焕然一新了。
#
#         """


handler = CrimeMining()
handler.main(content1)